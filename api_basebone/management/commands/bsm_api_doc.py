from docx import Document

from django.apps import apps
from django.core.management.base import BaseCommand
from api_basebone.export.fields import get_app_field_schema
from api_basebone.export.fields import get_model_field_config

from api_basebone.restful.client import views


def create_table_from_rows(document, l):
    """生成文档中的表格信息"""
    table = document.add_table(rows=len(l), cols=len(l[0]))
    table.style = 'Light List Accent 1'
    for i, row in enumerate(l):
        cells = table.rows[i].cells
        for j, s in enumerate(row):
            cells[j].text = s


def get_table_from_structure(document, app, model_name):
    l = []
    model_fullname = f'{app}__{model_name.lower()}'
    table_name = f'{app}_{model_name.lower()}'

    row = []
    row.append("属性")
    row.append("描述")
    row.append("必填")
    row.append("值类型")
    row.append("关联表")
    row.append("默认值")
    l.append(row)

    # app_config = apps.get_app_config(app)
    model_class = apps.get_model(app, model_name)
    field_config = get_model_field_config(model_class)
    fields = field_config[model_fullname].get('fields')
    for field in fields:
        row = []
        row.append(str(field.get('name', '')))
        row.append(str(field.get('displayName', '')))
        row.append(str(field.get('required', '')))
        row.append(str(field.get('type', '')))
        row.append(str(field.get('ref', '')))
        row.append(str(field.get('default', '')))
        l.append(row)

    create_table_from_rows(document, l)


def get_table_from_list(document, app, model_name, action):
    """"""
    l = []
    table_name = f'{app}_{model_name.lower()}'

    row = []
    row.append("属性")
    row.append("说明")
    l.append(row)

    row = []
    row.append("url")
    row.append(f"https://ip:port/basebone/client/{app}__{model_name.lower()}/list/")
    l.append(row)

    row = []
    row.append("访问方式")
    row.append("post")
    l.append(row)

    row = []
    row.append("作用")
    row.append(f"分页查询{table_name}表记录")
    l.append(row)

    row = []
    row.append("参数")
    row.append(
        """size：每页记录数(url参数)
page：页号(url参数)"""
    )
    l.append(row)

    row = []
    row.append("返回结果")
    row.append(
        f"""{{
    "error_code": "错误码，正常为0",
    "error_message": "错误信息，正常为空字符串",
    "result": [{{{table_name}}},{{{table_name}}}......] ////{table_name}表的分页信息，{{{table_name}}}的格式参照数据结构。
}}"""
    )
    l.append(row)

    create_table_from_rows(document, l)


def get_table_from_retrieve(document, app, model_name, action):
    """"""
    l = []
    table_name = f'{app}_{model_name.lower()}'

    row = []
    row.append("属性")
    row.append("说明")
    l.append(row)

    row = []
    row.append("url")
    row.append(
        f"""https://ip:port/basebone/client/{app}__{model_name.lower()}/<pk>/
<<pk>>为{table_name}表的主键"""
    )
    l.append(row)

    row = []
    row.append("访问方式")
    row.append("post")
    l.append(row)

    row = []
    row.append("作用")
    row.append(f"查询{table_name}表一条记录的明细")
    l.append(row)

    row = []
    row.append("参数")
    row.append(f"""除了URL中的主键，没有其他附加参数。""")
    l.append(row)

    row = []
    row.append("返回结果")
    row.append(
        f"""{{
    "error_code": "错误码，正常为0",
    "error_message": "错误信息，正常为空字符串",
    "result": {{{table_name}}}  ////{table_name}表的明细信息，{{{table_name}}}的格式参照数据结构。
}}"""
    )
    l.append(row)

    create_table_from_rows(document, l)


def get_table_from_update(document, app, model_name, action):
    """"""
    l = []
    table_name = f'{app}_{model_name.lower()}'

    row = []
    row.append("属性")
    row.append("说明")
    l.append(row)

    row = []
    row.append("url")
    row.append(
        f"""https://ip:port/basebone/client/{app}__{model_name.lower()}/<pk>/
<<pk>>为{table_name}表的主键"""
    )
    l.append(row)

    row = []
    row.append("访问方式")
    row.append("put")
    l.append(row)

    row = []
    row.append("作用")
    row.append(
        f"""根据主键更新{table_name}表一条记录。
{{{table_name}}}的格式参照数据结构。
跟partial_update操作不同的是，update需要提交完整的数据结构，必填的字段会作检测。"""
    )
    l.append(row)

    row = []
    row.append("参数")
    row.append(
        f"""{{{table_name}}}
////参数为body上传，格式为json，{{{table_name}}}的格式参照数据结构"""
    )
    l.append(row)

    row = []
    row.append("返回结果")
    row.append(
        f"""{{
    "error_code": "错误码，正常为0",
    "error_message": "错误信息，正常为空字符串",
    "result": {{{table_name}}}  ////{table_name}表的明细信息，{{{table_name}}}的格式参照数据结构。
}}"""
    )
    l.append(row)

    create_table_from_rows(document, l)


def get_table_from_partial_update(document, app, model_name, action):
    """"""
    l = []
    table_name = f'{app}_{model_name.lower()}'

    row = []
    row.append("属性")
    row.append("说明")
    l.append(row)

    row = []
    row.append("url")
    row.append(
        f"""https://ip:port/basebone/client/{app}__{model_name.lower()}/<pk>/
<<pk>>为{table_name}表的主键"""
    )
    l.append(row)

    row = []
    row.append("访问方式")
    row.append("patch")
    l.append(row)

    row = []
    row.append("作用")
    row.append(
        f"""根据主键更新{table_name}表一条记录。
跟update操作不同的是，partial_update不需要提交完整的数据结构，缺省的字段保持原来的值不变，必填的字段不作检测。"""
    )
    l.append(row)

    row = []
    row.append("参数")
    row.append(
        f"""{{{table_name}}}
////参数为body上传，格式为json，{{{table_name}}}的格式参照数据结构"""
    )
    l.append(row)

    row = []
    row.append("返回结果")
    row.append(
        f"""{{
    "error_code": "错误码，正常为0",
    "error_message": "错误信息，正常为空字符串",
    "result": {{{table_name}}}  ////{table_name}表的明细信息，{{{table_name}}}的格式参照数据结构
}}"""
    )
    l.append(row)

    create_table_from_rows(document, l)


def get_table_from_custom_patch(document, app, model_name, action):
    """"""
    l = []
    table_name = f'{app}_{model_name.lower()}'

    row = []
    row.append("属性")
    row.append("说明")
    l.append(row)

    row = []
    row.append("url")
    row.append(
        f"""https://ip:port/basebone/client/{app}__{model_name.lower()}/<pk>/patch/
<<pk>>为{table_name}表的主键"""
    )
    l.append(row)

    row = []
    row.append("访问方式")
    row.append("put")
    l.append(row)

    row = []
    row.append("作用")
    row.append(
        f"""根据主键更新{table_name}表一条记录。
跟update操作不同的是，custom_patch不需要提交完整的数据结构，缺省的字段保持原来的值不变，必填的字段不作检测。"""
    )
    l.append(row)

    row = []
    row.append("参数")
    row.append(
        f"""{{{table_name}}}
////参数为body上传，格式为json，{{{table_name}}}的格式参照数据结构"""
    )
    l.append(row)

    row = []
    row.append("返回结果")
    row.append(
        f"""{{
    "error_code": "错误码，正常为0",
    "error_message": "错误信息，正常为空字符串",
    "result": {{{table_name}}}  ////{table_name}表的明细信息，{{{table_name}}}的格式参照数据结构
}}"""
    )
    l.append(row)

    create_table_from_rows(document, l)


def get_table_from_create(document, app, model_name, action):
    """"""
    l = []
    table_name = f'{app}_{model_name.lower()}'

    row = []
    row.append("属性")
    row.append("说明")
    l.append(row)

    row = []
    row.append("url")
    row.append(f"""https://ip:port/basebone/client/{app}__{model_name.lower()}/""")
    l.append(row)

    row = []
    row.append("访问方式")
    row.append("post")
    l.append(row)

    row = []
    row.append("作用")
    row.append(
        f"""{table_name}表插入一条新记录。
{{{table_name}}}的格式参照数据结构。"""
    )
    l.append(row)

    row = []
    row.append("参数")
    row.append(
        f"""{{{table_name}}}
////参数为body上传，格式为json，{{{table_name}}}的格式参照数据结构"""
    )
    l.append(row)

    row = []
    row.append("返回结果")
    row.append(
        f"""{{
    "error_code": "错误码，正常为0",
    "error_message": "错误信息，正常为空字符串",
    "result": {{{table_name}}}  ////{table_name}表的明细信息，{{{table_name}}}的格式参照数据结构。
}}"""
    )
    l.append(row)

    create_table_from_rows(document, l)


ACTION_DOC_FUNCS = {
    "list": get_table_from_list,
    'retrieve': get_table_from_retrieve,
    # 'set': get_table_from_set,
    'update': get_table_from_update,
    'create': get_table_from_create,
    'partial_update': get_table_from_partial_update,
    'custom_patch': get_table_from_custom_patch,
}


class Command(BaseCommand):
    """输出模型配置

    只是简单的输出模型的配置，输出后的配置可进行调整和修改
    """

    def add_arguments(self, parser):
        """"""
        parser.add_argument('--app', type=str, help='指定导出api文档的app')

    def handle(self, *args, **kwargs):
        """"""
        self.stdout.write('export actions api document...')
        app = kwargs.get('app')

        document = Document()
        document.add_heading(f"app：{app}", 1)
        document.add_paragraph('')
        for key, data in views.exposed_apis.items():
            if not key.startswith(f'{app}__'):
                continue
            _, model_name = key.split('__')
            document.add_heading(f"模块：{model_name}", 2)
            document.add_heading(f"模块：{model_name}——数据结构", 3)
            get_table_from_structure(document, app, model_name)
            document.add_paragraph('')

            actions = set(data['actions'])
            if 'func' in actions:
                actions.remove('func')
            if 'set' in actions:
                actions.add('list')
                actions.remove('set')
            actions = list(actions)
            actions.sort()
            for action in actions:
                # if action == 'func':
                #     continue
                func = ACTION_DOC_FUNCS.get(action)
                if not func:
                    continue
                document.add_heading(f"模块：{model_name}——接口：{action}", 3)
                func(document, app, model_name, action)
                document.add_paragraph('')

            document.add_paragraph('')

        api_file = f'{app}.docx'
        try:
            print(f'-------------------开始导出 {app} api document ------------------')
            document.save(api_file)
            print(f'------------------- 导出 api document 结束 ----------------------')
        except Exception as e:
            print('导出 API document异常： {}'.format(str(e)))

